"""
tools/file_ops.py - File Operations Tool
Handles file creation, including .docx document generation.
"""

import os
from datetime import datetime
from langchain_core.tools import tool

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@tool
def create_docx(title: str, content: str, filename: str = None) -> str:
    """
    Create a Word document (.docx) with the given title and content.
    
    Args:
        title: The document title (will be the heading)
        content: The main content of the document (can include paragraphs separated by newlines)
        filename: Optional filename (without extension). If not provided, uses timestamp.
        
    Returns:
        The path to the created document
    """
    if not DOCX_AVAILABLE:
        return "Document creation requires python-docx. Install with: pip install python-docx"
    
    # Create output directory if it doesn't exist
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename if not provided
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"document_{timestamp}"
    
    filepath = os.path.join(output_dir, f"{filename}.docx")
    
    try:
        # Create document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content (split by double newlines for paragraphs)
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        # Add footer with timestamp
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save document
        doc.save(filepath)
        
        return f"✅ Document created: {filepath}"
        
    except Exception as e:
        return f"Error creating document: {str(e)}"


@tool
def save_text_file(content: str, filename: str, extension: str = "txt") -> str:
    """
    Save content to a text file.
    
    Args:
        content: The text content to save
        filename: The filename (without extension)
        extension: File extension (default: txt)
        
    Returns:
        The path to the created file
    """
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    
    filepath = os.path.join(output_dir, f"{filename}.{extension}")
    
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return f"✅ File saved: {filepath}"
    except Exception as e:
        return f"Error saving file: {str(e)}"
